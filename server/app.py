"""
Local LLM Chat UI — talks to Ollama on localhost:11434
Run: python app.py
"""

from __future__ import annotations

import asyncio
import base64
import io
import json
import re
import threading
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape

import httpx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from .markdown_export import Block, InlineRun, parse_blocks, parse_inline, soft_break_long_tokens
from .telemetry import sample as telemetry_sample
from .kill_switch import execute_kill_switch
from .serp_search import (
    WebSearchResult,
    build_system_prompt,
    classify_query,
    is_configured,
    load_serpapi_config,
    web_search,
)
from .weather import fetch_weather_line
from .reasoning_normalize import (
    ReasoningStreamNormalizer,
    apply_think_to_payload,
    looks_like_think_unsupported,
)

load_serpapi_config()

OLLAMA_BASE = "http://localhost:11434"
STATIC_DIR = Path(__file__).resolve().parent.parent / "static"
MAX_DOC_CHARS = 80_000

SYSTEM_PROMPT = (
    "You are a Local AI Software Engineer helping the user develop, debug, and maintain code.\n"
    "When coding, briefly outline logical steps (e.g. in a short comment block) before the final code.\n"
    "Write complete, working code — no placeholders or TODO stubs unless asked.\n"
    "Match existing project style; do not invent APIs, files, or schemas — ask or use attached context.\n"
    "If live context (time, weather) is provided below, treat it as ground truth and integrate it naturally.\n"
    "Never output bracketed citation markers such as [1], [2], or [3] unless web search sources require them.\n"
    "Never use LaTeX notation ($, $$, \\frac, \\sqrt, etc.). "
    "For math in prose use Unicode (\u00b2, \u221a, \u00d7, \u00f7, \u00b1). "
    "For math in code use standard operators (**, Math.sqrt, a/b).\n"
    "Omit greetings, apologies, and redundant summaries. Lead with the solution."
)

_TIME_INTENT_RE = re.compile(
    r"\b("
    r"what (day|time|date)|how long until|until |tonight|this morning|this evening|"
    r"weekday|timezone|what.?s the (time|date)|current time|today'?s date"
    r")\b",
    re.IGNORECASE,
)


def _needs_local_time(query: str) -> bool:
    return bool(_TIME_INTENT_RE.search(query or ""))


def _local_time_line() -> str:
    now = datetime.now().astimezone()
    return (
        f"Local time on this machine: {now.strftime('%A, %Y-%m-%d %H:%M %Z')} "
        f"(UTC{now.strftime('%z')}). Treat this as ground truth."
    )
IMAGE_TYPES = {"image/jpeg", "image/png", "image/gif", "image/webp", "image/bmp"}
DOC_TYPES = {
    "application/pdf",
    "text/plain",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

app = FastAPI(title="Local LLM Chat")
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


@app.get("/")
async def index():
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/api/telemetry")
async def telemetry():
    return telemetry_sample()


@app.post("/api/shutdown")
async def shutdown():
    threading.Thread(target=execute_kill_switch, daemon=True).start()
    return {"ok": True}


@app.get("/api/health")
async def health():
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            r = await client.get(f"{OLLAMA_BASE}/api/tags")
            r.raise_for_status()
            return {
                "ollama": True,
                "models": r.json().get("models", []),
                "web_search": is_configured(),
            }
    except Exception as e:
        return {"ollama": False, "error": str(e), "web_search": is_configured()}


@app.get("/api/web-search/status")
async def web_search_status():
    return {"configured": is_configured()}


VISION_NAME_RE = re.compile(
    r"vision|llava|bakllava|moondream|minicpm-v|pixtral|qwen",
    re.IGNORECASE,
)

THINKING_NAME_RE = re.compile(
    r"(^|[:/\-_])(r1|reason|think|qwq|deepseek-r1|openthinker|marco-o1|"
    r"gpt-oss|o1|o3|o4)([:/\-_]|$)",
    re.IGNORECASE,
)

# Digest -> capabilities cache so /api/models stays fast across polls.
_SHOW_CAP_CACHE: dict[str, list[str]] = {}


def _merge_capabilities(*sources: Any) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for src in sources:
        if not src:
            continue
        items = src if isinstance(src, (list, tuple, set)) else [src]
        for item in items:
            if not isinstance(item, str):
                continue
            key = item.strip().lower()
            if not key or key in seen:
                continue
            # Normalize common aliases.
            if key in ("reasoning", "reason", "think"):
                key = "thinking"
            seen.add(key)
            out.append(key)
    return out


def model_support(model: dict[str, Any]) -> dict[str, bool]:
    caps = set(model.get("capabilities") or [])
    name = model.get("name") or ""
    vision = "vision" in caps or bool(VISION_NAME_RE.search(name))
    # Capability metadata is advisory. Name heuristics catch older Ollama builds
    # that omit "thinking" from /api/tags even for reasoning models.
    thinking = "thinking" in caps or bool(THINKING_NAME_RE.search(name))
    return {
        "text": True,
        "images": vision,
        "pdf": True,
        "pdf_scanned": vision,
        "docx": True,
        "txt": True,
        "tools": "tools" in caps,
        "thinking": thinking,
    }


def model_labels(support: dict[str, bool]) -> list[dict[str, str]]:
    """Human-readable capability chips for the UI."""
    return [
        {
            "id": "text",
            "label": "Text",
            "status": "yes",
            "hint": "Plain text messages",
        },
        {
            "id": "images",
            "label": "Images",
            "status": "yes" if support["images"] else "no",
            "hint": "Photo analysis (JPG, PNG, WebP…)"
            if support["images"]
            else "Not a vision model — switch model to use images",
        },
        {
            "id": "pdf",
            "label": "PDF",
            "status": "yes" if support["pdf_scanned"] else "partial",
            "hint": "Text & scanned PDFs"
            if support["pdf_scanned"]
            else "Text PDFs only (scanned pages need a vision model)",
        },
        {
            "id": "docx",
            "label": "DOCX",
            "status": "yes",
            "hint": "Word documents (text extracted)",
        },
        {
            "id": "txt",
            "label": "TXT",
            "status": "yes",
            "hint": "Plain text files",
        },
        {
            "id": "thinking",
            "label": "Thinking",
            "status": "yes" if support["thinking"] else "partial",
            "hint": "Extended reasoning mode"
            if support["thinking"]
            else "Thought process panel always shown; model may not expose a trace",
        },
    ]


def model_context_length(model: dict[str, Any]) -> int:
    details = model.get("details") or {}
    ctx = details.get("context_length")
    if isinstance(ctx, int) and ctx > 0:
        return ctx
    family = (details.get("family") or "").lower()
    if "llama" in family:
        return 4096
    if "gemma" in family:
        return 8192
    return 8192


async def _show_capabilities(
    client: httpx.AsyncClient, name: str, digest: str
) -> list[str]:
    """Enrich capabilities via /api/show; cache by digest across polls."""
    cache_key = digest or name
    if cache_key in _SHOW_CAP_CACHE:
        return _SHOW_CAP_CACHE[cache_key]
    try:
        # Short per-model timeout so a slow show cannot kill /api/models.
        r = await client.post(
            f"{OLLAMA_BASE}/api/show",
            json={"name": name},
            timeout=3.0,
        )
        if r.status_code != 200:
            _SHOW_CAP_CACHE[cache_key] = []
            return []
        data = r.json()
        caps = _merge_capabilities(
            data.get("capabilities"),
            (data.get("model_info") or {}).get("capabilities"),
            (data.get("details") or {}).get("capabilities"),
        )
        # Some builds expose thinking under projector / family metadata only.
        family = str((data.get("details") or {}).get("family") or "").lower()
        if "think" in family or "reason" in family:
            caps = _merge_capabilities(caps, ["thinking"])
        _SHOW_CAP_CACHE[cache_key] = caps
        return caps
    except Exception:
        # Never fail model listing because show enrichment failed.
        return []


@app.get("/api/models")
async def list_models():
    try:
        async with httpx.AsyncClient(timeout=20.0) as client:
            r = await client.get(f"{OLLAMA_BASE}/api/tags")
            r.raise_for_status()
            models = r.json().get("models", [])
            # Parallel show enrichment — tags alone are enough if show fails.
            show_caps_list = await asyncio.gather(
                *[
                    _show_capabilities(
                        client, m.get("name") or "", m.get("digest") or ""
                    )
                    for m in models
                ]
            )
            enriched = []
            for m, show_caps in zip(models, show_caps_list):
                name = m.get("name") or ""
                digest = m.get("digest") or ""
                caps = _merge_capabilities(m.get("capabilities"), show_caps)
                if digest:
                    _SHOW_CAP_CACHE[digest] = list(caps)
                model_row = {
                    **m,
                    "capabilities": caps,
                }
                support = model_support(model_row)
                enriched.append(
                    {
                        "name": name,
                        "digest": digest,
                        "modified_at": m.get("modified_at") or m.get("modified") or "",
                        "size": m.get("size"),
                        "context_length": model_context_length(model_row),
                        "capabilities": caps,
                        "support": support,
                        "labels": model_labels(support),
                    }
                )
            enriched.sort(key=lambda row: row["name"].lower())
            # Drop stale cache entries for deleted digests.
            live = {row["digest"] or row["name"] for row in enriched}
            for key in list(_SHOW_CAP_CACHE):
                if key not in live:
                    _SHOW_CAP_CACHE.pop(key, None)
            return {"models": enriched}
    except httpx.ConnectError:
        raise HTTPException(
            503,
            "Cannot reach Ollama. Make sure Ollama is running (ollama serve).",
        )
    except Exception as e:
        raise HTTPException(500, str(e))


def extract_document_text(filename: str, content: bytes, mime: str) -> str:
    lower = filename.lower()

    if mime == "text/plain" or lower.endswith(".txt"):
        for enc in ("utf-8", "utf-16", "latin-1"):
            try:
                return content.decode(enc)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace")

    if mime == "application/pdf" or lower.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(content))
        parts = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
        return "\n\n".join(parts)

    if lower.endswith(".docx") or mime in DOC_TYPES:
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    raise ValueError(f"Unsupported document type: {filename}")


def build_user_message(
    text: str,
    images_b64: list[str],
    doc_text: str | None,
    doc_name: str | None,
) -> dict[str, Any]:
    parts = []
    if doc_text and doc_name:
        trimmed = doc_text[:MAX_DOC_CHARS]
        if len(doc_text) > MAX_DOC_CHARS:
            trimmed += "\n\n[Document truncated due to length limit.]"
        parts.append(
            f"--- Attached document: {doc_name} ---\n{trimmed}\n--- End of document ---"
        )
    if text.strip():
        parts.append(text.strip())
    elif images_b64:
        parts.append("Please analyze the attached image(s).")
    elif doc_text:
        parts.append("Please summarize and answer questions about the attached document.")

    content = "\n\n".join(parts)
    msg: dict[str, Any] = {"role": "user", "content": content}
    if images_b64:
        msg["images"] = images_b64
    return msg


def _parse_optional_float(raw: str | None) -> float | None:
    if raw is None:
        return None
    s = str(raw).strip()
    if not s:
        return None
    try:
        return float(s)
    except ValueError:
        return None


def _parse_optional_int(raw: str | None) -> int | None:
    if raw is None:
        return None
    s = str(raw).strip()
    if not s:
        return None
    try:
        return int(float(s))
    except ValueError:
        return None


def _merge_generation_options(
    options: dict[str, Any],
    *,
    temperature: str,
    top_p: str,
    top_k: str,
    num_predict: str,
    repeat_penalty: str,
) -> dict[str, Any]:
    temp = _parse_optional_float(temperature)
    if temp is not None and 0.0 <= temp <= 2.0:
        options["temperature"] = temp

    tp = _parse_optional_float(top_p)
    if tp is not None and 0.0 <= tp <= 1.0:
        options["top_p"] = tp

    tk = _parse_optional_int(top_k)
    if tk is not None and 1 <= tk <= 100:
        options["top_k"] = tk

    np_ = _parse_optional_int(num_predict)
    if np_ is not None and 1 <= np_ <= 131072:
        options["num_predict"] = np_

    rp = _parse_optional_float(repeat_penalty)
    if rp is not None and 0.5 <= rp <= 2.0:
        options["repeat_penalty"] = rp

    return options


@app.post("/api/chat")
async def chat(
    model: str = Form(...),
    messages: str = Form(...),
    stream: bool = Form(True),
    num_ctx: int = Form(4096),
    think: str = Form("auto"),
    web_search_enabled: str = Form("false"),
    lat: str = Form(""),
    lon: str = Form(""),
    elevation: str = Form(""),
    temperature: str = Form(""),
    top_p: str = Form(""),
    top_k: str = Form(""),
    num_predict: str = Form(""),
    repeat_penalty: str = Form(""),
    images: list[UploadFile] = File(default=[]),
    document: UploadFile | None = File(default=None),
):
    try:
        history: list[dict[str, Any]] = json.loads(messages)
    except json.JSONDecodeError:
        raise HTTPException(400, "Invalid messages JSON")

    images_b64: list[str] = []
    for img in images:
        if not img.filename:
            continue
        data = await img.read()
        if not data:
            continue
        mime = img.content_type or "application/octet-stream"
        if mime not in IMAGE_TYPES and not (img.filename or "").lower().endswith(
            (".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp")
        ):
            raise HTTPException(400, f"Unsupported image: {img.filename}")
        images_b64.append(base64.b64encode(data).decode("ascii"))

    doc_text: str | None = None
    doc_name: str | None = None
    if document and document.filename:
        doc_bytes = await document.read()
        if doc_bytes:
            doc_name = document.filename
            try:
                doc_text = extract_document_text(
                    document.filename,
                    doc_bytes,
                    document.content_type or "",
                )
            except Exception as e:
                raise HTTPException(400, f"Could not read document: {e}")
            if not doc_text.strip():
                raise HTTPException(400, "Document appears empty or unreadable.")

    # Attach files only to the latest user turn
    search_query = ""
    if history and history[-1].get("role") == "user":
        last = history[-1]
        search_query = (last.get("content") or "").strip()
        user_msg = build_user_message(
            last.get("content", ""),
            images_b64,
            doc_text,
            doc_name,
        )
        history[-1] = user_msg
    else:
        history.append(build_user_message("", images_b64, doc_text, doc_name))

    use_web_search = web_search_enabled.lower() in ("true", "1", "yes")
    if use_web_search and not is_configured():
        raise HTTPException(
            400,
            "Web search is enabled but SERPAPI_KEY is missing. "
            "Add your key to serpapikey.env or .env and restart the server.",
        )
    if use_web_search and not search_query:
        raise HTTPException(400, "Web search requires a text message.")

    search_type_hint = classify_query(search_query) if use_web_search else "general"
    coord_lat = _parse_optional_float(lat)
    coord_lon = _parse_optional_float(lon)
    coord_elev = _parse_optional_float(elevation)
    has_coords = coord_lat is not None and coord_lon is not None
    if has_coords and not (-90.0 <= coord_lat <= 90.0 and -180.0 <= coord_lon <= 180.0):
        has_coords = False
    if coord_elev is not None and not (-500.0 <= coord_elev <= 9000.0):
        coord_elev = None

    async def enrich_messages(
        msgs: list[dict[str, Any]],
    ) -> tuple[list[dict[str, Any]], WebSearchResult | None, bool]:
        parts: list[str] = [SYSTEM_PROMPT]
        weather_ok = False

        if has_coords:
            try:
                weather_line = await fetch_weather_line(
                    float(coord_lat),
                    float(coord_lon),
                    elevation=coord_elev,
                )
            except Exception:
                weather_line = None
            if weather_line:
                parts.append(weather_line)
                weather_ok = True

        if _needs_local_time(search_query):
            parts.append(_local_time_line())

        search_result: WebSearchResult | None = None
        if use_web_search:
            search_result = await web_search(search_query)
            parts.append(build_system_prompt(search_result))

        enriched = list(msgs)
        enriched.insert(0, {"role": "system", "content": "\n\n".join(parts)})
        return enriched, search_result, weather_ok

    ctx = max(512, min(num_ctx, 1_048_576))
    options: dict[str, Any] = {"num_ctx": ctx}
    _merge_generation_options(
        options,
        temperature=temperature,
        top_p=top_p,
        top_k=top_k,
        num_predict=num_predict,
        repeat_penalty=repeat_penalty,
    )
    base_payload: dict[str, Any] = {
        "model": model,
        "stream": stream,
        "options": options,
    }
    # Always request thinking when auto/true; retry without if unsupported.
    think_payload = apply_think_to_payload(base_payload, think)

    def _normalize_response(resp: dict[str, Any]) -> dict[str, Any]:
        normalizer = ReasoningStreamNormalizer()
        thinking, content, history_msg = normalizer.normalize_complete(resp)
        out = dict(resp)
        out["message"] = history_msg
        if thinking:
            out.setdefault("message", {})["thinking"] = thinking
        if content and not (out.get("message") or {}).get("content"):
            out["message"]["content"] = content
        return out

    if not stream:
        try:
            enriched_msgs, search_result, _ = await enrich_messages(history)
            async with httpx.AsyncClient(timeout=600.0) as client:
                payload = {**think_payload, "messages": enriched_msgs, "stream": False}
                r = await client.post(f"{OLLAMA_BASE}/api/chat", json=payload)
                body_text = r.text
                if r.status_code != 200:
                    if "think" in payload and looks_like_think_unsupported(body_text):
                        payload = {
                            **base_payload,
                            "messages": enriched_msgs,
                            "stream": False,
                        }
                        r = await client.post(f"{OLLAMA_BASE}/api/chat", json=payload)
                        body_text = r.text
                    r.raise_for_status()
                try:
                    raw = json.loads(body_text) if body_text else {}
                except json.JSONDecodeError as e:
                    raise HTTPException(502, f"Invalid JSON from Ollama: {e}") from e
                if isinstance(raw, dict) and raw.get("error"):
                    err = str(raw["error"])
                    if "think" in payload and looks_like_think_unsupported(err):
                        payload = {
                            **base_payload,
                            "messages": enriched_msgs,
                            "stream": False,
                        }
                        r = await client.post(f"{OLLAMA_BASE}/api/chat", json=payload)
                        r.raise_for_status()
                        raw = r.json()
                    else:
                        raise HTTPException(502, err)
                resp = _normalize_response(raw if isinstance(raw, dict) else {})
            if search_result is not None:
                resp["web_sources"] = search_result.sources
                resp["search_type"] = search_result.search_type
            return resp
        except HTTPException:
            raise
        except httpx.ConnectError:
            raise HTTPException(503, "Cannot reach Ollama.")
        except httpx.HTTPStatusError as e:
            raise HTTPException(e.response.status_code, e.response.text)
        except ValueError as e:
            raise HTTPException(400, str(e))
        except RuntimeError as e:
            raise HTTPException(502, str(e))

    async def event_stream():
        try:
            if has_coords:
                yield f"data: {json.dumps({'type': 'weather_start'})}\n\n"

            if use_web_search:
                yield f"data: {json.dumps({'type': 'search_start', 'query': search_query, 'search_type': search_type_hint})}\n\n"
                try:
                    yield f"data: {json.dumps({'type': 'search_fetch', 'count': 3})}\n\n"
                    ollama_history, search_result, weather_ok = await enrich_messages(history)
                except ValueError as e:
                    yield f"data: {json.dumps({'error': str(e)})}\n\n"
                    return
                except RuntimeError as e:
                    yield f"data: {json.dumps({'error': str(e)})}\n\n"
                    return
                if has_coords:
                    yield f"data: {json.dumps({'type': 'weather_done', 'ok': weather_ok})}\n\n"
                sources = search_result.sources if search_result else []
                stype = search_result.search_type if search_result else "general"
                yield f"data: {json.dumps({'type': 'search_done', 'query': search_query, 'search_type': stype, 'sources': sources})}\n\n"
            else:
                ollama_history, _, weather_ok = await enrich_messages(history)
                if has_coords:
                    yield f"data: {json.dumps({'type': 'weather_done', 'ok': weather_ok})}\n\n"

            payloads = [
                {**think_payload, "messages": ollama_history, "stream": True},
            ]
            if "think" in think_payload:
                payloads.append(
                    {**base_payload, "messages": ollama_history, "stream": True}
                )

            async with httpx.AsyncClient(timeout=600.0) as client:
                for attempt, payload in enumerate(payloads):
                    async with client.stream(
                        "POST", f"{OLLAMA_BASE}/api/chat", json=payload
                    ) as response:
                        if response.status_code != 200:
                            body = await response.aread()
                            err_text = body.decode(errors="replace")
                            if (
                                attempt == 0
                                and len(payloads) > 1
                                and looks_like_think_unsupported(err_text)
                            ):
                                continue
                            yield f"data: {json.dumps({'error': err_text})}\n\n"
                            return
                        normalizer = ReasoningStreamNormalizer()
                        emitted = False
                        retry_without_think = False
                        async for line in response.aiter_lines():
                            if not line:
                                continue
                            frames = normalizer.feed_line(line)
                            # Some Ollama builds return HTTP 200 with an error JSON body.
                            if (
                                not emitted
                                and attempt == 0
                                and len(payloads) > 1
                                and frames
                                and looks_like_think_unsupported(
                                    str(frames[0].get("error") or "")
                                )
                            ):
                                retry_without_think = True
                                break
                            for frame in frames:
                                emitted = True
                                yield f"data: {json.dumps(frame)}\n\n"
                        if retry_without_think:
                            continue
                        return
        except httpx.ConnectError:
            yield f"data: {json.dumps({'error': 'Cannot reach Ollama. Is it running?'})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


class ExportMessage(BaseModel):
    role: str
    content: str = ""
    timestamp: int | None = None
    modelName: str | None = None
    thinking: str | None = None
    sources: list[dict[str, Any]] | None = None
    docName: str | None = None


class ExportRequest(BaseModel):
    format: str
    title: str = "Chat export"
    messages: list[ExportMessage]


EXPORT_MEDIA_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


def export_slug(title: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", (title or "").lower()).strip("-")
    return (slug or "chat")[:60].strip("-") or "chat"


def export_filename(title: str, ext: str) -> str:
    return f"{export_slug(title)}-{datetime.now().strftime('%Y-%m-%d')}.{ext}"


def export_role_label(msg: ExportMessage) -> str:
    if msg.role == "user":
        return "You"
    if msg.role == "assistant":
        return msg.modelName or "Assistant"
    return msg.role.capitalize() if msg.role else "Message"


def export_timestamp(ts: int | None) -> str:
    if not ts:
        return ""
    try:
        return datetime.fromtimestamp(ts / 1000).strftime("%Y-%m-%d %H:%M")
    except (OverflowError, OSError, ValueError):
        return ""


def export_source_line(src: dict[str, Any]) -> str:
    idx = src.get("index")
    title = src.get("title") or src.get("url") or "Source"
    url = src.get("url") or ""
    prefix = f"[{idx}] " if idx is not None else ""
    return f"{prefix}{title} — {url}" if url else f"{prefix}{title}"


# ---------------------------------------------------------------------------
# PDF export — ReportLab
# ---------------------------------------------------------------------------

_WINDOWS_FONTS_DIR = Path(r"C:\Windows\Fonts")
FONT_REGULAR, FONT_BOLD, FONT_ITALIC, FONT_BOLDITALIC = (
    "Helvetica",
    "Helvetica-Bold",
    "Helvetica-Oblique",
    "Helvetica-BoldOblique",
)
FONT_MONO, FONT_MONO_BOLD = "Courier", "Courier-Bold"

try:
    pdfmetrics.registerFont(TTFont("ExportSans", str(_WINDOWS_FONTS_DIR / "segoeui.ttf")))
    pdfmetrics.registerFont(TTFont("ExportSans-Bold", str(_WINDOWS_FONTS_DIR / "segoeuib.ttf")))
    pdfmetrics.registerFont(TTFont("ExportSans-Italic", str(_WINDOWS_FONTS_DIR / "segoeuii.ttf")))
    pdfmetrics.registerFont(TTFont("ExportSans-BoldItalic", str(_WINDOWS_FONTS_DIR / "segoeuiz.ttf")))
    pdfmetrics.registerFontFamily(
        "ExportSans",
        normal="ExportSans",
        bold="ExportSans-Bold",
        italic="ExportSans-Italic",
        boldItalic="ExportSans-BoldItalic",
    )
    pdfmetrics.registerFont(TTFont("ExportMono", str(_WINDOWS_FONTS_DIR / "consola.ttf")))
    pdfmetrics.registerFont(TTFont("ExportMono-Bold", str(_WINDOWS_FONTS_DIR / "consolab.ttf")))
    FONT_REGULAR, FONT_BOLD, FONT_ITALIC, FONT_BOLDITALIC = (
        "ExportSans",
        "ExportSans-Bold",
        "ExportSans-Italic",
        "ExportSans-BoldItalic",
    )
    FONT_MONO, FONT_MONO_BOLD = "ExportMono", "ExportMono-Bold"
except Exception:
    pass  # fall back to the built-in base-14 PDF fonts (Latin-1 only)

PDF_PAGE_MARGIN = 0.9 * inch
PDF_CONTENT_WIDTH = LETTER[0] - 2 * PDF_PAGE_MARGIN


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    body = ParagraphStyle("ExportBody", parent=base["BodyText"], fontName=FONT_REGULAR, fontSize=10.5, leading=15)
    return {
        "Title": ParagraphStyle("ExportTitle", parent=base["Title"], fontName=FONT_BOLD),
        "Heading1": ParagraphStyle(
            "ExportH1", parent=base["Heading1"], fontName=FONT_BOLD, fontSize=15, leading=19, spaceAfter=8
        ),
        "Heading2": ParagraphStyle(
            "ExportH2", parent=base["Heading2"], fontName=FONT_BOLD, fontSize=13, leading=17, spaceAfter=6
        ),
        "Heading3": ParagraphStyle(
            "ExportH3", parent=base["Heading3"], fontName=FONT_BOLD, fontSize=11.5, leading=15, spaceAfter=6
        ),
        "Body": body,
        "Bullet": ParagraphStyle("ExportBullet", parent=body, leftIndent=18, spaceAfter=2),
        "Meta": ParagraphStyle(
            "ExportMeta", parent=body, fontName=FONT_BOLD, textColor=colors.HexColor("#333333"), spaceAfter=3
        ),
        "Label": ParagraphStyle(
            "ExportLabel",
            parent=body,
            fontName=FONT_BOLDITALIC,
            textColor=colors.HexColor("#555555"),
            spaceBefore=6,
            spaceAfter=2,
        ),
        "Code": ParagraphStyle(
            "ExportCode",
            parent=body,
            fontName=FONT_MONO,
            fontSize=9,
            leading=12,
            backColor=colors.HexColor("#f2f2f2"),
            borderPadding=6,
            spaceAfter=6,
        ),
    }


def _pdf_inline_markup(text: str) -> str:
    parts = []
    for run in parse_inline(text):
        content = xml_escape(run.text).replace("\n", "<br/>")
        if run.code:
            content = f'<font face="{FONT_MONO}">{content}</font>'
        if run.bold:
            content = f"<b>{content}</b>"
        if run.italic:
            content = f"<i>{content}</i>"
        parts.append(content)
    return "".join(parts)


def _pdf_table(rows: list[list[str]], styles: dict[str, ParagraphStyle]) -> Table:
    n_cols = max((len(r) for r in rows), default=1) or 1
    col_width = PDF_CONTENT_WIDTH / n_cols
    data = []
    for r_idx, row in enumerate(rows):
        style = styles["Meta"] if r_idx == 0 else styles["Body"]
        cells = [Paragraph(_pdf_inline_markup(row[c]) or "&nbsp;", style) for c in range(len(row))]
        while len(cells) < n_cols:
            cells.append(Paragraph("&nbsp;", style))
        data.append(cells)
    table = Table(data, colWidths=[col_width] * n_cols)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _render_blocks_pdf(blocks: list[Block], styles: dict[str, ParagraphStyle], flow: list[Any]) -> None:
    number_counter = 0
    prev_kind: str | None = None
    for block in blocks:
        if block.kind == "number":
            number_counter = number_counter + 1 if prev_kind == "number" else 1
        else:
            number_counter = 0

        if block.kind == "heading":
            level = min(max(block.level, 1), 3)
            flow.append(Paragraph(_pdf_inline_markup(block.text), styles[f"Heading{level}"]))
        elif block.kind == "paragraph":
            flow.append(Paragraph(_pdf_inline_markup(block.text) or "&nbsp;", styles["Body"]))
            flow.append(Spacer(1, 4))
        elif block.kind == "bullet":
            flow.append(Paragraph(f"\u2022&nbsp;&nbsp;{_pdf_inline_markup(block.text)}", styles["Bullet"]))
        elif block.kind == "number":
            flow.append(Paragraph(f"{number_counter}.&nbsp;&nbsp;{_pdf_inline_markup(block.text)}", styles["Bullet"]))
        elif block.kind == "code":
            code_text = "\n".join(soft_break_long_tokens(line, max_run=60) for line in block.lines)
            flow.append(Preformatted(code_text, styles["Code"]))
        elif block.kind == "table" and block.rows:
            flow.append(_pdf_table(block.rows, styles))
            flow.append(Spacer(1, 6))
        elif block.kind == "hr":
            flow.append(HRFlowable(width="100%", color=colors.HexColor("#bbbbbb"), thickness=0.75, spaceAfter=4))
        prev_kind = block.kind


def build_pdf_export(title: str, messages: list[ExportMessage]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=PDF_PAGE_MARGIN,
        rightMargin=PDF_PAGE_MARGIN,
        topMargin=PDF_PAGE_MARGIN,
        bottomMargin=PDF_PAGE_MARGIN,
    )
    styles = _pdf_styles()
    flow: list[Any] = [
        Paragraph(xml_escape(title or "Chat export"), styles["Title"]),
        Spacer(1, 14),
    ]
    for msg in messages:
        ts = export_timestamp(msg.timestamp)
        label = export_role_label(msg)
        header = f"{label} — {ts}" if ts else label
        flow.append(Paragraph(xml_escape(header), styles["Meta"]))
        _render_blocks_pdf(parse_blocks(msg.content or ""), styles, flow)

        if msg.thinking:
            flow.append(Paragraph("Thinking", styles["Label"]))
            _render_blocks_pdf(parse_blocks(msg.thinking), styles, flow)

        if msg.sources:
            flow.append(Paragraph("Sources", styles["Label"]))
            for src in msg.sources:
                text = soft_break_long_tokens(export_source_line(src))
                flow.append(Paragraph(xml_escape(text), styles["Body"]))

        if msg.docName:
            flow.append(Paragraph(f"Attached: {xml_escape(msg.docName)}", styles["Body"]))

        flow.append(Spacer(1, 10))
    doc.build(flow)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX export — python-docx
# ---------------------------------------------------------------------------


def _docx_add_inline_runs(paragraph, text: str) -> None:
    runs: list[InlineRun] = parse_inline(text)
    if not runs:
        paragraph.add_run("")
        return
    for run_data in runs:
        run = paragraph.add_run(run_data.text)
        run.bold = run_data.bold
        run.italic = run_data.italic
        if run_data.code:
            run.font.name = "Consolas"


def _docx_add_paragraph_lines(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    lines = text.split("\n") or [""]
    for idx, line in enumerate(lines):
        _docx_add_inline_runs(p, line)
        if idx < len(lines) - 1:
            p.add_run().add_break()
    return p


def _docx_add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell_p = table.cell(r_idx, c_idx).paragraphs[0]
            _docx_add_inline_runs(cell_p, cell_text)
            if r_idx == 0:
                for run in cell_p.runs:
                    run.bold = True


def _docx_add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _render_blocks_docx(doc: Document, blocks: list[Block]) -> None:
    for block in blocks:
        if block.kind == "heading":
            level = min(max(block.level, 1), 4)
            _docx_add_paragraph_lines(doc, block.text, style=f"Heading {level}")
        elif block.kind == "paragraph":
            _docx_add_paragraph_lines(doc, block.text)
        elif block.kind == "bullet":
            _docx_add_paragraph_lines(doc, block.text, style="List Bullet")
        elif block.kind == "number":
            _docx_add_paragraph_lines(doc, block.text, style="List Number")
        elif block.kind == "code":
            code_text = "\n".join(soft_break_long_tokens(line, max_run=60) for line in block.lines)
            p = doc.add_paragraph()
            code_lines = code_text.split("\n")
            for idx, line in enumerate(code_lines):
                run = p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                if idx < len(code_lines) - 1:
                    p.add_run().add_break()
        elif block.kind == "table":
            _docx_add_table(doc, block.rows)
        elif block.kind == "hr":
            _docx_add_horizontal_rule(doc)


def build_docx_export(title: str, messages: list[ExportMessage]) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Segoe UI"
    doc.add_heading(title or "Chat export", level=1)

    for msg in messages:
        ts = export_timestamp(msg.timestamp)
        label = export_role_label(msg)
        header = f"{label} — {ts}" if ts else label
        p = doc.add_paragraph()
        run = p.add_run(header)
        run.bold = True

        _render_blocks_docx(doc, parse_blocks(msg.content or ""))

        if msg.thinking:
            label_p = doc.add_paragraph()
            label_run = label_p.add_run("Thinking")
            label_run.bold = True
            label_run.italic = True
            _render_blocks_docx(doc, parse_blocks(msg.thinking))

        if msg.sources:
            label_p = doc.add_paragraph()
            label_run = label_p.add_run("Sources")
            label_run.bold = True
            for src in msg.sources:
                doc.add_paragraph(export_source_line(src), style="List Number")

        if msg.docName:
            doc.add_paragraph(f"Attached: {msg.docName}")

        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.post("/api/export")
async def export_chat(req: ExportRequest):
    fmt = req.format.lower()
    if fmt not in EXPORT_MEDIA_TYPES:
        raise HTTPException(400, "Unsupported export format")
    if not req.messages:
        raise HTTPException(400, "Chat has no messages to export")

    try:
        if fmt == "docx":
            data = build_docx_export(req.title, req.messages)
        else:
            data = build_pdf_export(req.title, req.messages)
    except Exception as e:
        raise HTTPException(500, f"Could not build export: {e}")

    filename = export_filename(req.title, fmt)
    return Response(
        content=data,
        media_type=EXPORT_MEDIA_TYPES[fmt],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


if __name__ == "__main__":
    import uvicorn

    print("\n  Local LLM Chat UI")
    print("  Open http://127.0.0.1:7860 in your browser\n")
    uvicorn.run(app, host="127.0.0.1", port=7860, log_level="info")
